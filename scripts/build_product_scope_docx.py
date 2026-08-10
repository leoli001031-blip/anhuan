from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Cm, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "环保托管运营平台_产品定义与一期范围_2026-08-04.md"
OUTPUT = ROOT / "output" / "doc" / "环保托管运营平台_产品定义与一期范围_2026-08-04.docx"

TEAL = "0F5C5E"
TEAL_DARK = "0A4143"
TEAL_LIGHT = "E7F2F1"
INK = "24323D"
MUTED = "5E6D78"
LINE = "C8D9D8"
AMBER = "E5A33B"
WHITE = "FFFFFF"
FONT_CN = "Source Han Sans CN"
FONT_LATIN = "Source Han Sans CN"
FONT_MONO = "Source Code Pro"


def set_run_font(run, size=None, bold=None, color=None, name=FONT_CN):
    run.font.name = name
    r_fonts = run._element.get_or_add_rPr().rFonts
    for script in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{script}"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(element, fill):
    if hasattr(element, "_p"):
        element = element._p
    elif hasattr(element, "_tc"):
        element = element._tc
    props = element.get_or_add_tcPr() if element.tag.endswith("tc") else element.get_or_add_pPr()
    shd = props.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        props.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_hyperlink(paragraph, label, url, size=None):
    rel_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT_LATIN)
    r_fonts.set(qn("w:eastAsia"), FONT_CN)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), TEAL)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend((r_fonts, color, underline))
    if size is not None:
        for tag_name in ("sz", "szCs"):
            size_node = OxmlElement(f"w:{tag_name}")
            size_node.set(qn("w:val"), str(int(size * 2)))
            r_pr.append(size_node)
    text = OxmlElement("w:t")
    text.text = label
    run.extend((r_pr, text))
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


TOKEN_RE = re.compile(r"(\[[^\]]+\]\([^)]+\)|\*\*[^*]+\*\*|`[^`]+`)")


def add_inline(paragraph, text, base_size=None, base_color=INK):
    for token in TOKEN_RE.split(text.strip()):
        if not token:
            continue
        link = re.fullmatch(r"\[([^\]]+)\]\(([^)]+)\)", token)
        if link:
            add_hyperlink(paragraph, link.group(1), link.group(2), base_size)
        elif token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, base_size, True, base_color)
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            font = FONT_CN if re.search(r"[\u3400-\u9fff]", token) else FONT_MONO
            set_run_font(run, base_size, False, TEAL_DARK, font)
            shade(paragraph, "F4F7F7")
        else:
            run = paragraph.add_run(token)
            set_run_font(run, base_size, False, base_color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_borders(table, color=LINE, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def add_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.autofit = True
    set_table_borders(table)
    for r_idx, row_values in enumerate(rows):
        row = table.rows[r_idx]
        prevent_row_split(row)
        if r_idx == 0:
            set_repeat_table_header(row)
        for c_idx, value in enumerate(row_values):
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                shade(cell._tc, TEAL)
            elif r_idx % 2 == 0:
                shade(cell._tc, "F4F8F8")
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            add_inline(paragraph, value, 8.2, WHITE if r_idx == 0 else INK)
            if r_idx == 0:
                for run in paragraph.runs:
                    run.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_callout(doc, text, fill=TEAL_LIGHT, accent=TEAL):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Mm(3)
    table.columns[1].width = Mm(163)
    set_table_borders(table, fill, "0")
    left, body = table.rows[0].cells
    shade(left._tc, accent)
    shade(body._tc, fill)
    set_cell_margins(left, 80, 0, 80, 0)
    set_cell_margins(body, 170, 180, 170, 180)
    p = body.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.28
    add_inline(p, text, 10.5, INK)
    return table


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, 8, False, MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, end))
    tail = paragraph.add_run(" 页")
    set_run_font(tail, 8, False, MUTED)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_CN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.25

    heading1 = styles["Heading 1"]
    heading1.font.name = FONT_CN
    heading1._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    heading1.font.size = Pt(17)
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor.from_string(TEAL_DARK)
    heading1.paragraph_format.space_before = Pt(16)
    heading1.paragraph_format.space_after = Pt(8)
    heading1.paragraph_format.keep_with_next = True

    heading2 = styles["Heading 2"]
    heading2.font.name = FONT_CN
    heading2._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    heading2.font.size = Pt(12.5)
    heading2.font.bold = True
    heading2.font.color.rgb = RGBColor.from_string(TEAL)
    heading2.paragraph_format.space_before = Pt(10)
    heading2.paragraph_format.space_after = Pt(5)
    heading2.paragraph_format.keep_with_next = True

    for name, left, hanging in (("Scope Bullet", 0.65, 0.35), ("Scope Number", 0.72, 0.5)):
        if name not in styles:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = styles[name]
        style.base_style = normal
        style.paragraph_format.left_indent = Cm(left)
        style.paragraph_format.first_line_indent = Cm(-hanging)
        style.paragraph_format.space_after = Pt(3)


def configure_section(doc):
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(21)
    section.right_margin = Mm(21)
    section.header_distance = Mm(8)
    section.footer_distance = Mm(8)
    section.different_first_page_header_footer = True

    header_p = section.header.paragraphs[0]
    header_p.text = "环保管家业务操作系统  /  产品定义与一期范围"
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header_p.runs:
        set_run_font(run, 7.5, False, MUTED)

    add_page_number(section.footer.paragraphs[0])
    first_footer = section.first_page_footer.paragraphs[0]
    first_footer.text = "SCOPE DRAFT  ·  2026-08-04"
    first_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in first_footer.runs:
        set_run_font(run, 8, False, MUTED)


def build_cover(doc):
    for _ in range(3):
        doc.add_paragraph()

    accent = doc.add_table(rows=1, cols=1)
    accent.rows[0].height = Mm(3)
    shade(accent.cell(0, 0)._tc, TEAL)
    set_table_borders(accent, TEAL, "0")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("环保管家业务操作系统")
    set_run_font(run, 28, True, TEAL_DARK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(22)
    run = p.add_run("产品定义与一期范围")
    set_run_font(run, 19, False, TEAL)

    add_callout(
        doc,
        "核心定位：不是通用知识库，而是把多企业环保资料、结构化台账、证据检索、整改任务、服务时间线、客户门户与经营跟进连成闭环的多租户运营平台。",
    )

    doc.add_paragraph()
    meta = doc.add_table(rows=4, cols=2)
    meta.autofit = False
    set_table_borders(meta, LINE, "4")
    meta_data = [
        ("版本", "v0.2 研究与范围稿"),
        ("状态", "SCOPE_DRAFT · 待首个试点企业确认"),
        ("一期建议", "1 家服务商 / 1 个地区 / 1 个主场景 / 3–5 家企业 / 网页端"),
        ("日期", "2026-08-04"),
    ]
    for idx, (label, value) in enumerate(meta_data):
        left, right = meta.rows[idx].cells
        set_cell_margins(left, 120, 150, 120, 150)
        set_cell_margins(right, 120, 150, 120, 150)
        shade(left._tc, "F0F6F5")
        for cell, text_value, is_label in ((left, label, True), (right, value, False)):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text_value)
            set_run_font(run, 9.5, is_label, TEAL_DARK if is_label else INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("设计边界")
    set_run_font(run, 9, True, AMBER)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_inline(
        p,
        "本文用于把口述需求、旧黑客松 Demo 审计和公开行业资料固化为开发基线；法规与数据合规内容不替代律师、网安机构或环保专业人员的正式意见。",
        9,
        MUTED,
    )
    p.add_run().add_break(WD_BREAK.PAGE)


def parse_markdown(doc, text):
    lines = text.splitlines()
    start = next(i for i, line in enumerate(lines) if line.startswith("## 1."))
    i = start
    current_h2 = ""
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("```"):
            language = stripped[3:].strip()
            block = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            if language == "mermaid":
                p = doc.add_paragraph(style="Heading 2")
                p.add_run("业务关系（文字版）")
                add_table(
                    doc,
                    [
                        ["入口 / 资料源", "进入平台后的结果"],
                        ["环保服务商运营后台", "跨企业服务总览、分工、质量与经营跟进"],
                        ["企业客户门户", "仅本企业资料、问题、待办、报告与反馈"],
                        ["合作伙伴工作台", "仅授权期内被分配企业的允许内容"],
                        ["企业档案 + 共享法规", "权限过滤后检索，形成带页码证据的发现项"],
                        ["发现项", "整改任务、服务日历、专家复核与版本化报告"],
                        ["人工确认的发现项", "可另行转为服务机会，不改变专业结论"],
                    ],
                )
            else:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(block))
                set_run_font(run, 8.5, False, TEAL_DARK, FONT_MONO)
                shade(p, "F4F7F7")
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            parsed = []
            for table_line in table_lines:
                values = [value.strip() for value in table_line.strip("|").split("|")]
                if all(re.fullmatch(r":?-{3,}:?", value) for value in values):
                    continue
                parsed.append(values)
            if parsed:
                add_table(doc, parsed)
            continue

        heading = re.match(r"^(##|###)\s+(.+)$", stripped)
        if heading:
            level, title = heading.groups()
            if level == "##":
                current_h2 = title
            p = doc.add_paragraph(style="Heading 1" if level == "##" else "Heading 2")
            add_inline(p, title)
            i += 1
            continue

        if stripped.startswith("> "):
            add_callout(doc, stripped[2:])
            i += 1
            continue

        bullet = re.match(r"^-\s+(.+)$", stripped)
        if bullet:
            p = doc.add_paragraph(style="Scope Bullet")
            source_list = current_h2.startswith("15. 公开资料来源")
            if source_list:
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.line_spacing = 1.05
            font_size = 8.4 if source_list else 10.3
            run = p.add_run("• ")
            set_run_font(run, font_size, True, TEAL)
            add_inline(p, bullet.group(1), font_size)
            i += 1
            continue

        numbered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if numbered:
            p = doc.add_paragraph(style="Scope Number")
            run = p.add_run(f"{numbered.group(1)}. ")
            set_run_font(run, 10.3, True, TEAL)
            add_inline(p, numbered.group(2), 10.3)
            i += 1
            continue

        p = doc.add_paragraph()
        add_inline(p, stripped, 10.5)
        i += 1


def main():
    source_text = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    doc.core_properties.title = "环保管家业务操作系统：产品定义与一期范围"
    doc.core_properties.subject = "多租户环保托管平台产品范围与一期开发基线"
    doc.core_properties.author = "Codex × 项目团队"
    doc.core_properties.keywords = "环保管家, 一企一档, OCR, 多租户, 合规运营"

    configure_styles(doc)
    configure_section(doc)
    build_cover(doc)
    parse_markdown(doc, source_text)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
